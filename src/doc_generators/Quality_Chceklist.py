from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx2pdf import convert
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import random
import os
import datetime
from docx.shared import Inches
from docx.enum.table import WD_ALIGN_VERTICAL

#---lists used by the generator---#
customers = ['NORWAY', 'POLAND', 'CANADA', 'BRAZIL', 'GREECE', 'FRANCE', 'TURKEY', 'SWEDEN', 'BELGIUM', 'FINLAND']
products = [
    'MC-540X', 'TR-200B', 'HF-390A', 'PL-601Z', 'DX-777T',
    'TX-820V', 'MX-450L', 'RX-310Z', 'VF-220D', 'GL-980S',
    'AL-115Q', 'KP-320E', 'BZ-660F', 'QN-770H', 'SL-430M',
    'ZR-205R', 'TY-350G', 'XK-610U', 'JD-700W', 'CN-150C',
    'VR-940T', 'MS-600P', 'LK-890B', 'FT-730X', 'NE-245A',
    'PW-515Y', 'RM-860N', 'WD-180S', 'KV-390K', 'CE-905L',
    'LP-555V', 'GH-770J', 'SB-140D', 'QP-660F', 'NU-440Z',
    'AZ-300T', 'XD-710R', 'RE-850C', 'MR-160H', 'TL-900X'
]
items = [
    'Steel Sheet A36', 'Hex Bolts M12', 'Rubber Gasket 80mm', 'Bearing 6202 ZZ', 'Shaft 500mm',
    'Packaging Box L', 'Copper Wire 3mm', 'Insulation Foam Pad', 'Control Panel Mount',
    'Plastic Cover 150x150', 'Aluminum Bracket', 'Heat Resistant Sleeve', 'Stainless Bolt M8',
    'Clamp Ring 120mm', 'Sensor Clip', 'Ceramic Disc 80mm', 'Rubber Stopper', 'Spacer 2mm',
    'Terminal Block 4P', 'Ventilation Grid', 'Plastic Rivets', 'Spring Washer M10',
    'Grease Tube 250ml', 'Epoxy Resin Kit', 'Protective Sleeve 50mm', 'Digital Display Unit',
    'Cable Tie Pack (100)', 'Gasket Sheet A4', 'O-Ring NBR 60mm', 'Fuse 5A', 'LED Light Strip',
    'Cooling Gel Pack', 'Graphite Pad', 'Support Foot Steel', 'Wiring Loom 1m',
    'Insulated Tube 25mm', 'Pressure Valve', 'Capacitor 450V', 'Thermal Fuse',
    'Nut M6', 'Connector 2P', 'Heat Sink ALU', 'Teflon Tape Roll', 'Battery Pack','Hinge Set','Power Switch','Wooden Pallet'
]

fonts = ['Arial', 'Calibri', 'Aptos'] 
checklist_points = [
    "Shipping mark is illegible or missing",
    "Carton is damaged or markings incorrect",
    "Incorrect quantity or assortment",
    "Wrong product size",
    "Packaging does not match signed sample",
    "Package is not sealed completely",
    "Missing distributor information",
    "Missing logo or warning label",
    "Instruction manual is missing or damaged",
    "Wood splinter or sharp point on product",
    "Exposed nail with sharp point",
    "Dead or live insect in packaging",
    "Rubber texture or glossiness mismatch",
    "Paint smearing or scratches",
    "Loose parts inside packaging",
    "Sharp edges on plastic components",
    "Product doesn't power on",
    "Incorrect barcode or label",
    "Connector not working",
    "Incorrect orientation in box",
    "Visual defect on housing",
    "Screws loose or missing",
    "Functionality test failed",
    "Not assembled as per drawing",
    "Missing safety labels",
    "Color mismatch",
    "Battery not included",
    "Rubber Switch not working",
    "Hinges loose",
    "Dust/debris inside packaging"
]

inspectors = [
    'Anna Nowak', 'Jan Kowalski', 'Piotr Lewandowski', 'Agnieszka Zielińska',
    'Magdalena Witkowska', 'Nadia Sauter', 'Bartosz Wawrzyniak', 'Mateusz Jarzyna',
    'Wiktor Kopczyński', 'Przemysław Wąsik', 'Dawid Oszmiańczuk']

signature_roles = [
    ("Supervised by:","Date:"),
    ("Checked by:","Timestamp:"),
    ("Verified by:","On:")
]

column_synonyms = {
    "Product ID":      ["Product ID","Item ID","Part No."],
    "Checklist No.":   ["Checklist No.","Form ID","QC Ref."],
    "Inspection Point":["Inspection Point","Check Item","Step"],
    "Sampling Level":  ["Sampling Level","AQL Tier","Audit Level"],
    "Classification":  ["Classification","Class","Category"],
    "CR":              ["CR","Critical"],
    "MA":              ["MA","Major"],
    "MI":              ["MI","Minor"],
    "Hold":            ["Hold","Quarantine"],
    "Notes":           ["Notes","Comments","Remarks"]
}

qc_intros = [
    "This checklist captures quality inspection points and sampling levels.",
    "Below are the items to be verified during the final product review.",
    "The following table outlines inspection criteria and classification levels.",
    "Please review each checklist point and mark the sampling results.",
    "This section details quality requirements and test points for the batch.",
    "Use this list to confirm adherence to AQL and safety standards.",
    "Entries include both visual and functional inspection items.",
    "Ensure all non-conforming marks are clearly documented.",
    "Refer to the quality register for sampling-plan references.",
    "This summary supports the production-release quality gate.",
    "Check that inspection steps follow the approved procedure.",
    "Use this extract to coordinate sign-off with the QA manager.",
    "All checklist entries are timestamped for traceability.",
    "Confirm that sampling levels comply with customer agreements.",
    "Archive this list in the quality-management system.",
    "This closure summary indicates compliance with inspection criteria."
]

qc_summaries = [
    "All critical and major inspection points have been addressed.",
    "Items marked for hold require additional review before release.",
    "Refer to notes for any observed defects or deviations.",
    "Overall quality status indicates compliance with defined AQL levels.",
    "Please ensure supervised sign-off on any non-conforming points.",
    "Checklist results have been reported to the quality manager.",
    "Corrective actions are scheduled for identified issues.",
    "Inspection summary is filed for regulatory compliance.",
    "Confirm that all sampling results are within acceptable limits.",
    "This summary supports the end-of-line quality certification.",
    "Flag any open issues in the CAPA tracking system.",
    "Archive this summary in the audit-readiness folder.",
    "Use this closure report to update the quality KPI dashboard.",
    "Ensure that all remarks have corresponding evidence attachments.",
    "This final note confirms the checklist is complete and approved.",
    "All summary comments have been validated by the QA team."
]

def get_synonym(hdr):
    return random.choice(column_synonyms.get(hdr, [hdr]))

def add_horizontal_line(paragraph):
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'auto')
    pBdr.append(bottom)
    pPr.append(pBdr)

def add_quality_header(doc, layout_type):

    titles = [
        "Quality Checklist", "Inspection Sheet", "Conformance Log",
        "Audit Trail", "Verification Log", "Inspection Checklist",
        "Quality Inspection List", "Inspection Register", "QC Checklist",
        "Compliance Log", "Quality Review", "Inspection Record",
        "Quality Verification", "Inspection Summary", "Conformance Report"
    ]
    
    title_text = random.choice(titles)
    
    para = doc.add_paragraph()
    para.paragraph_format.tab_stops.add_tab_stop(Inches(6.0), alignment=WD_PARAGRAPH_ALIGNMENT.RIGHT)
    run = para.add_run(title_text)
    run.font.size = Pt(11)
    run.bold = True

    if layout_type == 1:
        run.add_text(f"\t#{random.randint(1000000, 9999999)}")
        add_horizontal_line(para)
        date_para = doc.add_paragraph()
        date_run = date_para.add_run(f"Date: {datetime.date.today().strftime('%d.%m.%Y')}")
        date_run.font.size = Pt(11)
        date_run.bold = True

    elif layout_type == 2:
        add_horizontal_line(para)

def add_product_info_table(doc):
    table = doc.add_table(rows=2, cols=4)
    table.style = 'Table Grid'

    run_00 = table.cell(0, 0).paragraphs[0].add_run("Product ID")
    run_00.bold = True
    table.cell(0, 1).text = random.choice(products)

    run_02 = table.cell(0, 2).paragraphs[0].add_run("Customer ID")
    run_02.bold = True
    table.cell(0, 3).text = random.choice(customers)

    run_10 = table.cell(1, 0).paragraphs[0].add_run("Item Description")
    run_10.bold = True

    merged_cell = table.cell(1, 1).merge(table.cell(1, 2)).merge(table.cell(1, 3))
    merged_cell.text = random.choice(items)
    doc.add_paragraph("") 

def add_aql_table(doc):
    table = doc.add_table(rows=3, cols=4)
    table.style = 'Table Grid'

    run_aql = table.cell(0, 0).paragraphs[0].add_run("AQL Level")
    run_aql.bold = True
    table.cell(0, 1).text = "Critical"
    table.cell(0, 2).text = "Major"
    table.cell(0, 3).text = "Minor"

    run_default = table.cell(1, 0).paragraphs[0].add_run("Default")
    run_default.bold = True
    table.cell(1, 1).text = "0"
    table.cell(1, 2).text = "2.5"
    table.cell(1, 3).text = "4.0"

    run_cust = table.cell(2, 0).paragraphs[0].add_run("Customer specific")
    run_cust.bold = True
    v_col = random.choice([1, 2, 3])
    for i in range(1, 4):
        table.cell(2, i).text = "V" if i == v_col else ""

def add_supervision_section(doc):
    para = doc.add_paragraph()
    run = para.add_run("Supervised by: __________________      Date: ______________")
    run.bold = True
    run.font.size = Pt(9)

def generate_checklist_table(doc):
    inspection_syns      = ["Inspection Checklist Points", "Quality Check Items", "Review Points", "Audit Criteria"]
    sampling_level_syns  = ["Sampling Level", "Inspection Depth", "Check Intensity", "Sample Tier"]
    classification_syns  = ["Classification", "Category", "Defect Class", "Severity"]
    notes_syns           = ["Notes", "Comments", "Remarks", "Observations"]
    product_req_syns     = ["Product Requirements", "Item Specs", "Component Criteria", "Design Specs"]
    header_letter_syns   = [["A","B","C"], ["X","Y","Z"], ["1","2","3"]]
    class_labels_syns    = [["CR","MA","MI","Hold"], ["C","M","m","H"], ["Critical","Major","Minor","OnHold"]]

    num_rows = random.randint(5, 18)
    selected_points = random.sample(checklist_points, num_rows)

    table = doc.add_table(rows=2 + num_rows, cols=8)
    table.style = 'Table Grid'

    col_widths = [
        Inches(0.4),   # No
        Inches(2.6),   # Inspection Points
        Inches(1.0),   # Sampling Level
        Inches(0.5),   # Classification (merged)
        Inches(0.5),   # CR / first class
        Inches(0.5),   # MA / second class
        Inches(0.5),   # MI / third class
        Inches(2.0),   # Notes
    ]
    for row in table.rows:
        for i, w in enumerate(col_widths):
            row.cells[i].width = w

    hdr1 = table.rows[0].cells
    hdr1[0].text = ""
    hdr1[1].paragraphs[0].add_run(random.choice(inspection_syns)).bold = True
    hdr1[2].paragraphs[0].add_run(random.choice(sampling_level_syns)).bold = True
    hdr1[3].paragraphs[0].add_run(random.choice(classification_syns)).bold = True
    hdr1[3].merge(table.cell(0, 6))
    hdr1[7].paragraphs[0].add_run(random.choice(notes_syns)).bold = True

    hdr2 = table.rows[1].cells
    hdr2[0].text = random.choice(header_letter_syns)[random.randrange(3)]
    hdr2[1].text = random.choice(product_req_syns)
    hdr2[2].text = ""
    for i, lbl in enumerate(random.choice(class_labels_syns)):
        hdr2[i+3].paragraphs[0].add_run(lbl).bold = True
    hdr2[7].text = ""

    for idx, point in enumerate(selected_points):
        row = table.rows[idx+2].cells
        row[0].text = str(idx+1)
        row[1].text = point
        row[2].text = random.choice(["", "Level I", "Level II", "Level III"])

        marked = random.sample(range(3, 7), k=random.randint(0, 2))
        for i in range(3, 7):
            row[i].text = "V" if i in marked else ""

        row[7].text = random.choice(["", "Needs review", "Critical impact", "Minor issue observed"])

        for cell in row:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

def add_reference_standards(doc):
    doc.add_paragraph("Reference Standards:")
    table = doc.add_table(rows=3, cols=2)
    table.style = 'Table Grid'
    table.cell(0,0).text = "Standard"
    table.cell(0,1).text = "Edition"
    table.cell(1,0).text = random.choice(["ISO 9001", "CE Directive", "RoHS"])
    table.cell(1,1).text = random.choice(["2015", "2020", "2011"])
    table.cell(2,0).text = random.choice(["IEC 61010", "UL 61010"])
    table.cell(2,1).text = random.choice(["3rd Ed.", "4th Ed."])
    doc.add_paragraph("")

def generate_quality_checklist(doc_number: int, output_folder: str):
    doc = Document()
    layout_type = random.randint(1, 3)
    #layout_type  = 3
    font_name = random.choice(fonts)
    font_size = Pt(9)
    style = doc.styles['Normal']
    style.font.name = font_name
    style.font.size = font_size

    add_quality_header(doc, layout_type)

    if layout_type == 1:
        add_product_info_table(doc)

        intro_sentences = random.sample(qc_intros, k=random.randint(5,10))
        if intro_sentences:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after  = Pt(0)
            for i, sent in enumerate(intro_sentences):
                run = p.add_run(sent)
                if i < len(intro_sentences) - 1:
                    run.add_text(" ")
        doc.add_paragraph("")

        generate_checklist_table(doc)
        doc.add_paragraph("")
        add_aql_table(doc)

        doc.add_paragraph("")
        if random.random() < 0.6:
            intro_sentences = random.sample(qc_summaries, k=random.randint(4,8))
            if intro_sentences:
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after  = Pt(0)
                for i, sent in enumerate(intro_sentences):
                    run = p.add_run(sent)
                    if i < len(intro_sentences) - 1:
                        run.add_text(" ")

    elif layout_type == 2:
        intro_sentences = random.sample(qc_intros, k=random.randint(4,8))
        if intro_sentences:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after  = Pt(0)
            for i, sent in enumerate(intro_sentences):
                run = p.add_run(sent)
                if i < len(intro_sentences) - 1:
                    run.add_text(" ")

        doc.add_paragraph("")
        add_product_info_table(doc)
        generate_checklist_table(doc)
        doc.add_paragraph("")
        add_supervision_section(doc) 

    elif layout_type == 3:
        generate_checklist_table(doc)
        add_reference_standards(doc)

    os.makedirs(output_folder, exist_ok=True)
    docx_path = os.path.join(output_folder, f"quality_checklist_{doc_number:03}_{layout_type}.docx")
    pdf_path  = os.path.join(output_folder, f"quality_checklist_{doc_number:03}_{layout_type}.pdf")
    doc.save(docx_path)
    convert(docx_path, pdf_path)
    os.remove(docx_path)

