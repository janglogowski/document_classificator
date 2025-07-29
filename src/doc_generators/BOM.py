import os
import random
import datetime
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx2pdf import convert
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

#---lists used by the generator---#
customers = ['NORWAY', 'POLAND', 'CANADA', 'BRAZIL', 'GREECE', 'FRANCE', 'TURKEY', 'SWEDEN', 'BELGIUM', 'FINLAND']
supervisors = ['Anna Nowak', 'Jan Kowalski', 'Peter Schmidt', 'Laura Rossi', 'Carlos Garcia']
products = [
    'MC-540X', 'TR-200B', 'HF-390A', 'PL-601Z', 'DX-777T',
    'TX-820V', 'MX-450L', 'RX-310Z', 'VF-220D', 'GL-980S',
    'AL-115Q', 'KP-320E', 'BZ-660F', 'QN-770H', 'SL-430M',
    'ZR-205R', 'TY-350G', 'XK-610U', 'JD-700W', 'CN-150C',
    'VR-940T', 'MS-600P', 'LK-890B', 'FT-730X', 'NE-245A'
]
units = ['kg', 'pcs', 'm', 'liters']
fonts = ['Arial', 'Calibri', 'Aptos']
items = [
    ('Steel Sheet A36', 'kg', 5.00), ('Hex Bolts M12', 'pcs', 0.25), ('Rubber Gasket 80mm', 'pcs', 0.50),
    ('Bearing 6202 ZZ', 'pcs', 1.50), ('Shaft 500mm', 'pcs', 8.00), ('Packaging Box L', 'pcs', 1.00),
    ('Copper Wire 3mm', 'm', 0.60), ('Insulation Foam Pad', 'pcs', 3.20), ('Control Panel Mount', 'pcs', 12.00),
    ('Plastic Cover 150x150', 'pcs', 1.10), ('Aluminum Bracket', 'pcs', 4.50), ('Heat Resistant Sleeve', 'm', 2.70),
    ('Stainless Bolt M8', 'pcs', 0.35), ('Clamp Ring 120mm', 'pcs', 1.75), ('Sensor Clip', 'pcs', 0.95),
    ('Ceramic Disc 80mm', 'pcs', 2.10), ('Rubber Stopper', 'pcs', 0.55), ('Spacer 2mm', 'pcs', 0.15),
    ('Terminal Block 4P', 'pcs', 3.40), ('Ventilation Grid', 'pcs', 5.60), ('Plastic Rivets', 'pcs', 0.20),
    ('Spring Washer M10', 'pcs', 0.05), ('Grease Tube 250ml', 'pcs', 1.90), ('Epoxy Resin Kit', 'pcs', 7.30),
    ('Protective Sleeve 50mm', 'm', 1.80), ('Digital Display Unit', 'pcs', 15.00), ('Cable Tie Pack (100)', 'pcs', 0.95),
    ('Gasket Sheet A4', 'pcs', 1.25), ('Fuse 5A', 'pcs', 0.30), ('LED Light Strip', 'm', 2.50),
    ('Battery Pack', 'pcs', 25.00), ('Hinge Set', 'pcs', 2.50), ('Power Switch', 'pcs', 1.20),
    ('Wooden Pallet', 'pcs', 15.00)
]
remarks_list = [
    '', '', 'SKF brand', 'High grade', 'Certified batch', 'Eco compliant',
    'Imported', 'ROHS compliant', 'ISO-verified', 'Urgent', 'For export',
    'Li-Ion battery installed', 'Hinge alignment adjusted', 'Switch tested OK'
]
bom_titles = ["BILL OF MATERIAL", "COMPONENT BREAKDOWN", "BOM Report"]
header_synonyms = {
    "Item Description": ["Component", "Part Name", "Element"],
    "Qty": ["Quantity", "QTY", "Total"],
    "Unit Price": ["Rate", "Cost/unit"],
    "Amount": ["Line Total", "Sum", "Aggregate"],
    "Remarks": ["Notes", "Comments"]
}

bom_intros = [
    "This document provides a detailed breakdown of all components required for the assembly process.",
    "Below is the component listing and associated costs for the upcoming production batch.",
    "The following table summarizes the materials and quantities needed for the current project.",
    "Please review the itemized list of parts and material specifications before procurement.",
    "This section outlines the parts, unit prices and total amounts for assembly.",
    "Use this breakdown to verify sourcing and cost estimates.",
    "All entries reflect the latest inventory and supplier rates.",
    "Ensure each component meets the specified quality standards.",
    "Refer to this parts register to plan raw-material purchasing.",
    "The component roster below includes unit costs and batch codes.",
    "This summary lists every item required, with per-unit pricing details.",
    "Review the materials tally for compliance with budget allowances.",
    "The parts manifest here is designed to support procurement workflows.",
    "All line-item costs are current as per vendor quotes.",
    "This extract shows the bill of components and total projected spend.",
    "Use this schedule of parts to align with sourcing and stock levels."
]

bom_summaries = [
    "All listed components have been verified for availability and compliance.",
    "Totals include estimated over-consumption allowances and current unit rates.",
    "Please confirm supplier lead times to ensure timely delivery of all items.",
    "Amounts reflect current pricing; adjust as necessary for bulk orders.",
    "Verify that all remark items meet the sourcing department’s standards.",
    "Final amounts include handling and logistics costs where applicable.",
    "Review this summary against the master budgeting sheet.",
    "All sourcing notes have been logged for audit purposes.",
    "Ensure this materials summary is reconciled with the purchase order.",
    "The cost subtotal supports financial forecasting for the next cycle.",
    "Check that component quantities align with production run requirements.",
    "This final review confirms that all items are ready for requisition.",
    "Cross-verify totals with the ERP system for consistency.",
    "Any deviations from standard pricing have been annotated here.",
    "This closure summary validates that all parts are approved for release.",
    "Ensure archival of this materials summary for compliance records."
]

def add_horizontal_line(paragraph):
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'),'single')
    bottom.set(qn('w:sz'),'4')
    bottom.set(qn('w:space'),'1')
    bottom.set(qn('w:color'),'auto')
    pBdr.append(bottom)
    pPr.append(pBdr)

def get_synonym(header):
    return random.choice(header_synonyms.get(header,[header]))

def add_customer_info_table(doc, customer, supervisor, date, product_id, internal_number, order_qty):
    table = doc.add_table(rows=2, cols=3)
    table.style='Table Grid'; table.autofit=True
    table.rows[0].cells[0].text=f"Customer ID: {customer}"
    table.rows[0].cells[1].text=f"Coordinator: {supervisor}"
    table.rows[0].cells[2].text=f"Date: {date.strftime('%Y-%m-%d')}"
    table.rows[1].cells[0].text=f"Product ID: {product_id}"
    table.rows[1].cells[1].text=f"Internal No.: {internal_number}"
    table.rows[1].cells[2].text=f"Order Qty: {order_qty}"

def add_total_amount_line(doc, total):
    line = doc.add_paragraph()
    add_horizontal_line(line)
    para = doc.add_paragraph()
    run = para.add_run(f"TOTAL:  {total:,.2f}")
    run.bold = True

def add_total_amount_table(doc, total):
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    table.rows[0].cells[0].text = "Total Amount: "
    table.rows[0].cells[1].text = "{:,.2f}".format(total)
    doc.add_paragraph("")

def generate_bom_table(doc, layout_type):
    if layout_type == 2:
        num_items = random.randint(3, 6)
        headers = ["No", "Item Description", "Qty", "UOM", "Unit Price", "Amount", "Remarks"] 

        table = doc.add_table(rows=len(headers), cols=num_items + 1)
        table.style = 'Table Grid'

        for i, header in enumerate(headers):
            table.cell(i, 0).text = header
            for p in table.cell(i, 0).paragraphs:
                for r in p.runs:
                    r.bold = True

        total = 0
        for col in range(1, num_items + 1):
            desc, uom, unit_price = random.choice(items)
            base_qty = random.randint(100, 999)
            qty = base_qty
            amount = round(qty * unit_price, 2)
            total += amount

            values = [
                str(col),
                desc,
                str(qty),
                uom,
                f"{unit_price:.2f}",
                f"{amount:,.2f}",
                random.choice(remarks_list)
            ]

            for row_idx, val in enumerate(values):
                table.cell(row_idx, col).text = val

        return total

    elif layout_type == 3:
        table = doc.add_table(rows=1, cols=7)
        headers = ["No", "Item Description", "Qty", "UOM", "Rate", "Amount", "Remarks"]
    else:
        table = doc.add_table(rows=1, cols=9)
        headers = ["No", "Item Description", "Consumption", "Extra %", "Qty", "UOM", "Rate", "Amount", "Remarks"]

    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        if layout_type == 1:
            for paragraph in hdr_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.bold = True

    total = 0
    if layout_type == 1:
        num_items = random.randint(4, 10)
    elif layout_type == 3:
        num_items = random.randint(8, 12)

    for i in range(num_items):
        desc, uom, unit_price = random.choice(items)
        consumption = round(random.uniform(0.3, 3.5), 2)
        extra_pct = random.choice([0, 2, 5, 10])
        base_qty = random.randint(1, 50)
        qty = int(base_qty * (1 + extra_pct / 100))
        amount = round(qty * unit_price, 2)
        total += amount

        if layout_type == 3:
            row = table.add_row().cells
            row[0].text = str(i + 1)
            row[1].text = desc
            row[2].text = str(qty)
            row[3].text = uom
            row[4].text = f"{unit_price:.2f}"
            row[5].text = f"{amount:,.2f}"
            row[6].text = random.choice(remarks_list)

        elif layout_type == 1:
            row = table.add_row().cells
            row[0].text = str(i + 1)
            row[1].text = desc
            row[2].text = str(consumption)
            row[3].text = f"{extra_pct}%"
            row[4].text = str(qty)
            row[5].text = uom
            row[6].text = f"{unit_price:.2f}"
            row[7].text = f"{amount:,.2f}"
            row[8].text = random.choice(remarks_list)
    return total

def generate_bom_document(doc_number: int, output_folder: str):
    doc = Document()
    layout_type = random.randint(1, 3)
    #layout_type = 3
    font_name = random.choice(fonts)
    font_size = Pt(10)
    style = doc.styles['Normal']
    style.font.name = font_name
    style.font.size = font_size

    customer = random.choice(customers)
    supervisor = random.choice(supervisors)
    product_id = random.choice(products)
    order_qty = random.randint(50, 500)
    internal_number = random.randint(1000000, 9999999)
    today = datetime.date.today() - datetime.timedelta(days=random.randint(0, 1000))
    title_text = random.choice(bom_titles)

    if layout_type == 1:
        heading = doc.add_paragraph(title_text)
        heading.runs[0].bold = True
        add_horizontal_line(heading)
        add_customer_info_table(doc, customer, supervisor, today, product_id, internal_number, order_qty)
        doc.add_paragraph("")

        if random.random() < 0.7:
            intro_sentences = random.sample(bom_intros, k=random.randint(1,4))
            if intro_sentences:
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after  = Pt(0)
                for i, sent in enumerate(intro_sentences):
                    run = p.add_run(sent)
                    if i < len(intro_sentences) - 1:
                        run.add_text(" ")
            doc.add_paragraph("")

        total = generate_bom_table(doc, layout_type)
        doc.add_paragraph("")

        if random.random() < 0.7:
            summary_sentences = random.sample(bom_summaries, k=random.randint(1,4))
            if summary_sentences:
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after  = Pt(0)
                for i, sent in enumerate(summary_sentences):
                    run = p.add_run(sent)
                    if i < len(summary_sentences) - 1:
                        run.add_text(" ")
            doc.add_paragraph("")

        add_total_amount_table(doc, total)
        doc.add_paragraph("\nApproved By: _________________          Sourcing Department: _________________")

    elif layout_type == 2:
        para1 = doc.add_paragraph()
        para1.paragraph_format.tab_stops.add_tab_stop(Inches(6.0), alignment=WD_PARAGRAPH_ALIGNMENT.RIGHT)
        run1 = para1.add_run(f"Customer ID: {customer}\tPrepared by: {supervisor}")
        run1.font.name = font_name
        run1.font.size = Pt(12)
        run1.bold = True

        para2 = doc.add_paragraph()
        para2.paragraph_format.tab_stops.add_tab_stop(Inches(6.0), alignment=WD_PARAGRAPH_ALIGNMENT.RIGHT)
        run2 = para2.add_run(f"Product ID: {product_id}\tInternal No.: {internal_number}")
        run2.font.name = font_name
        run2.font.size = Pt(12)
        run2.bold = True

        if random.random() < 0.6:
            line = doc.add_paragraph()
            add_horizontal_line(line)
            intro_sentences = random.sample(bom_intros, k=random.randint(4,6))
            if intro_sentences:
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after  = Pt(0)
                for i, sent in enumerate(intro_sentences):
                    run = p.add_run(sent)
                    if i < len(intro_sentences) - 1:
                        run.add_text(" ")
            doc.add_paragraph("")

        total = generate_bom_table(doc, layout_type)
        add_total_amount_line(doc, total)

    elif layout_type == 3:
        heading = doc.add_paragraph(title_text)

        if random.random() < 0.7:
            line = doc.add_paragraph()
            add_horizontal_line(line)
            intro_sentences = random.sample(bom_intros, k=random.randint(4,6))
            if intro_sentences:
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after  = Pt(0)
                for i, sent in enumerate(intro_sentences):
                    run = p.add_run(sent)
                    if i < len(intro_sentences) - 1:
                        run.add_text(" ")
            doc.add_paragraph("") 

        total = generate_bom_table(doc, layout_type)
        doc.add_paragraph("") 

        if random.random() < 0.7:
            summary_sentences = random.sample(bom_summaries, k=random.randint(1,4))
            if summary_sentences:
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after  = Pt(0)
                for i, sent in enumerate(summary_sentences):
                    run = p.add_run(sent)
                    if i < len(summary_sentences) - 1:
                        run.add_text(" ")

        line = doc.add_paragraph()
        add_horizontal_line(line)
        add_customer_info_table(doc, customer, supervisor, today, product_id, internal_number, order_qty)
        doc.add_paragraph("") 
        add_total_amount_table(doc, total)

    #---save docx, convert to pdf, remove docx---#
    os.makedirs(output_folder, exist_ok=True)
    docx_path = os.path.join(output_folder, f"bom_{doc_number:03}_{layout_type}.docx")
    pdf_path  = os.path.join(output_folder, f"bom_{doc_number:03}_{layout_type}.pdf")
    doc.save(docx_path)
    convert(docx_path, pdf_path)
    os.remove(docx_path)