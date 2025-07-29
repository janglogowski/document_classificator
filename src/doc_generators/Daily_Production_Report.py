from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx2pdf import convert
import random
import os
import datetime

#---lists used by the generator---#
customers = ['NORWAY', 'POLAND', 'CANADA', 'BRAZIL', 'GREECE', 'FRANCE', 'TURKEY', 'SWEDEN', 'BELGIUM', 'FINLAND']
fonts = ['Arial', 'Calibri', 'Aptos']
supervisors = ['Anna Nowak', 'Jan Kowalski', 'Peter Schmidt', 'Laura Rossi', 'Carlos Garcia']
operators = [
    'Anna Nowak', 'Jan Kowalski', 'Piotr Lewandowski', 'Agnieszka Zielińska',
    'Magdalena Witkowska', 'Nadia Sauter', 'Bartosz Wawrzyniak', 'Mateusz Jarzyna',
    'Wiktor Kopczyński', 'Przemysław Wąsik', 'Dawid Oszmiańczuk']
departments = ["Machining", "Assembly", "Quality Control", "Painting", "Packaging", "Logistics", "Maintenance"]
shifts = ["A (Morning)", "B (Evening)", "C (Night)"]
machine_ids = [f"MC-2{str(i).zfill(2)}" for i in range(1, 10)]
products = [
    'MC-540X','TR-200B','HF-390A','PL-601Z','DX-777T','TX-820V','MX-450L','RX-310Z','VF-220D','GL-980S',
    'AL-115Q','KP-320E','BZ-660F','QN-770H','SL-430M','ZR-205R','TY-350G','XK-610U','JD-700W','CN-150C',
    'VR-940T','MS-600P','LK-890B','FT-730X','NE-245A','PW-515Y','RM-860N','WD-180S','KV-390K','CE-905L',
    'LP-555V','GH-770J','SB-140D','QP-660F','NU-440Z','AZ-300T','XD-710R','RE-850C','MR-160H','TL-900X'
]

column_synonyms = {
    "Product ID":     ["Product ID","Item ID","Part ID"],
    "Target Qty":     ["Target Qty","Planned Qty","Planned Output"],
    "Actual Qty":     ["Actual Qty","Achieved Qty","Produced Qty"],
    "Scrap Qty":      ["Scrap Qty","Rejected Qty","Defect Qty"],
    "Scrap %":        ["Scrap %","Rejection %","Failure %"],
    "Rework Qty":     ["Rework Qty","Rework Count","Reworked Units"],
    "Remarks":        ["Remarks","Notes","Comments","Uwagi"],
    "Machine ID":     ["Machine ID","Machine","Equipment"],
    "Operation":      ["Operation","Process","Task"],
    "Operator":       ["Operator","Worker","Technician"],
    "Start Time":     ["Start Time","Begin","From"],
    "End Time":       ["End Time","Finish","To"],
    "Duration (min)": ["Duration (min)","Time Spent","Total Time","Total duration"],
    "Temperature":    ["Temperature","Temp (°C)","Thermal Reading"],
    "Energy (kWh)":   ["Energy (kWh)","Energy Used","Power Draw"],
    "Status":         ["Status","Stage","Condition"]
}

daily_intros = [
    "This report summarizes the daily operations and output for the production line.",
    "Below is an overview of machine performance and output figures for today’s shift.",
    "The following data captures production targets, actual output, and any deviations.",
    "Please review today’s operation log and output summary for each workstation.",
    "This section details the key metrics and activities recorded during the shift.",
    "Use this overview to assess throughput and efficiency levels.",
    "Entries include all start/stop times and operator assignments.",
    "Ensure any downtime incidents are noted for follow-up.",
    "Refer to this log to monitor shift productivity and cycle counts.",
    "The summary below lists each station’s actual versus planned output.",
    "This extract shows daily yield, scrap rates, and rework occurrences.",
    "The line-output register captures all production events for the day.",
    "Review the operational data to track compliance with daily goals.",
    "This dashboard section highlights any production bottlenecks.",
    "Use this shift summary to drive continuous improvement actions.",
    "All throughput figures are recorded for capacity planning."
]

daily_summaries = [
    "All production targets have been logged; deviations are highlighted above.",
    "No critical delays were observed; please address any minor issues noted.",
    "Ensure routine maintenance between shifts to maintain consistent output.",
    "Refer to remarks for any rework or quality concerns.",
    "Overall production performance met expectations for the day.",
    "Record any adjustments to shift schedules or staffing here.",
    "Verify the final counts against inventory records.",
    "All operator notes have been archived for review.",
    "Confirm that scrap percentages align with quality benchmarks.",
    "This closure summary signals readiness for the next production run.",
    "Note any unscheduled stops in the downtime register.",
    "Check material consumption against standard usage rates.",
    "Ensure shift-handover notes include any pending issues.",
    "The performance recap supports the morning briefing agenda.",
    "Archive this output summary for end-of-day reporting.",
    "Use this summary to update the overall production dashboard."
]

def get_synonym(col):
    return random.choice(column_synonyms.get(col, [col]))

def add_horizontal_line(paragraph):
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'auto')
    pBdr.append(bottom)
    pPr.append(pBdr)

def add_daily_report_title(doc, layout_type):
    if random.random() < 0.2:
        return

    title = doc.add_paragraph()
    title.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT if layout_type == 2 else WD_PARAGRAPH_ALIGNMENT.LEFT
    title_text = random.choice(["Production Log", "DAILY PRODUCTION REPORT", "Shift Summary"])
    run = title.add_run(title_text)
    run.font.size = Pt(14)
    run.bold = True

    if layout_type in [1, 2]:
        add_horizontal_line(title)

def add_daily_report_info_table(doc, layout_type):
    date_str = (datetime.date.today() - datetime.timedelta(days=random.randint(0, 730))).strftime('%d-%m-%Y')
    customer = f"Customer: {random.choice(customers)}"
    report_suffix = f"{random.randint(100,999):03}"
    report_no_full = f"Report No.: PR-{report_suffix}" if layout_type == 3 else f"PR-{report_suffix}"
    shift = f"Shift: {random.choice(shifts)}"

    if layout_type == 3:
        para1 = doc.add_paragraph()
        para1.paragraph_format.tab_stops.add_tab_stop(Inches(3.0), WD_PARAGRAPH_ALIGNMENT.LEFT)
        run1 = para1.add_run(f"{report_no_full}\t{customer}")
        run1.bold = True
        para2 = doc.add_paragraph()
        para2.paragraph_format.tab_stops.add_tab_stop(Inches(3.0), WD_PARAGRAPH_ALIGNMENT.LEFT)
        run2 = para2.add_run(f"{shift}\tDate: {date_str}")
        run2.bold = True

    elif layout_type == 2:
        table = doc.add_table(rows=5, cols=2)
        table.style = 'Table Grid'
        labels = ["Report No.:", "Customer:", "Supervisor:", "Department:", "Shift:"]
        values = [
            report_no_full,
            random.choice(customers),
            random.choice(supervisors),
            random.choice(departments),
            random.choice(shifts)
        ]
        for i in range(5):
            table.cell(i, 0).text = labels[i]
            table.cell(i, 1).text = values[i]

        date_p = doc.add_paragraph(date_str)
        date_p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        date_p.paragraph_format.space_before = Pt(0)
        date_p.paragraph_format.space_after = Pt(8)

    else:
        table = doc.add_table(rows=6, cols=2)
        table.style = 'Table Grid'
        labels = ["Report No.:", "Customer:", "Supervisor:", "Department:", "Date:", "Shift:"]
        values = [
            report_no_full,
            random.choice(customers),
            random.choice(supervisors),
            random.choice(departments),
            date_str,
            random.choice(shifts)
        ]
        for i in range(6):
            table.cell(i, 0).text = labels[i]
            table.cell(i, 1).text = values[i]
        doc.add_paragraph("")

def add_machine_operation_table(doc: Document, layout_type: int):
    base = ["Machine ID","Operation","Operator","Start Time","End Time","Duration (min)","Remarks"]
    if layout_type==2:
        base += ["Temperature","Energy (kWh)"]
    elif layout_type==3:
        base += ["Status"]
    syn = [get_synonym(c) for c in base]
    tbl = doc.add_table(rows=1,cols=len(syn), style='Table Grid')
    for i,h in enumerate(syn):
        cell = tbl.cell(0,i); cell.text = h
        cell.paragraphs[0].runs[0].bold = True

    ops = ["Welding","Drilling","Cutting","Assembly","Polishing"]
    for _ in range(random.randint(5,12)):
        r = tbl.add_row().cells
        start  = random.randint(6,14)
        startm = random.choice([0,15,30]); dur=random.randint(20,120)
        endm   = (start*60+startm+dur)%60; end=(start*60+startm+dur)//60
        vals = [
            random.choice(machine_ids),
            random.choice(ops),
            random.choice(operators),
            f"{start:02}:{startm:02}",
            f"{end:02}:{endm:02}",
            str(dur),
            random.choice(["OK","N/A","Delay","Recalibrated",""])
        ]
        if layout_type==2:
            vals += [f"{random.uniform(55,80):.1f}",f"{random.uniform(1,3):.2f}"]
        elif layout_type==3:
            vals += [random.choice(["In progress","Completed","Delayed"])]
        for i,v in enumerate(vals):
            r[i].text = v
    doc.add_paragraph()

def add_production_output_summary(doc, layout_type):
    if layout_type == 2:
        return
    cols = 7 if layout_type == 3 else 6
    table = doc.add_table(rows=1, cols=cols)
    table.style = 'Table Grid'

    headers = ["Product ID", "Target Qty", "Actual Qty", "Scrap Qty", "Scrap %"]
    if layout_type == 3:
        headers.append("Rework Qty")
    headers.append("Remarks")

    syn_headers = [get_synonym(h) for h in headers]
    for i, h in enumerate(syn_headers):
        table.cell(0, i).text = h
        table.cell(0, i).paragraphs[0].runs[0].bold = True

    for _ in range(random.randint(3, 6)):
        target = random.randint(100, 300)
        actual = random.randint(int(target * 0.85), target)
        scrap = target - actual
        scrap_pct = round((scrap / target) * 100, 2)
        values = [
            random.choice(products),
            str(target),
            str(actual),
            str(scrap),
            f"{scrap_pct:.2f}%" 
        ]
        if layout_type == 3:
            values.append(str(random.randint(0, 10)))
        values.append(random.choice(["", "Rework needed", "Scrap confirmed"]))
        row = table.add_row().cells
        for i, val in enumerate(values):
            row[i].text = val
    doc.add_paragraph("")

def add_signature_section(doc, layout_type):
    if layout_type == 3:
        para = doc.add_paragraph("Approved by: ___________        Prepared by: ___________")
        para.paragraph_format.space_before = Pt(20)
    else:
        table = doc.add_table(rows=2, cols=2)
        table.style = 'Table Grid'
        roles = ["Prepared by:", "Approved by:"]
        people = [random.choice(supervisors), random.choice(supervisors)]
        for i in range(2):
            table.cell(i, 0).text = roles[i]
            table.cell(i, 1).text = people[i]

def generate_daily_production_report(doc_number: int, output_folder: str):
    doc = Document()
    layout_type = random.randint(1,3)
    #layout_type = 3
    style = doc.styles['Normal']
    style.font.name = random.choice(fonts)
    style.font.size = Pt(9)

    add_daily_report_title(doc, layout_type)
    add_daily_report_info_table(doc, layout_type)
    if layout_type == 2:
        if random.random() < 0.7:
            intro_sentences = random.sample(daily_intros, k=random.randint(3,4))
            if intro_sentences:
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after  = Pt(0)
                for i, sent in enumerate(intro_sentences):
                    run = p.add_run(sent)
                    if i < len(intro_sentences) - 1:
                        run.add_text(" ")
            doc.add_paragraph("")
    add_machine_operation_table(doc, layout_type)

    if layout_type == 3:
        if random.random() < 0.7:
            intro_sentences = random.sample(daily_summaries, k=random.randint(4,6))
            if intro_sentences:
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after  = Pt(0)
                for i, sent in enumerate(intro_sentences):
                    run = p.add_run(sent)
                    if i < len(intro_sentences) - 1:
                        run.add_text(" ")
            doc.add_paragraph("")

    add_production_output_summary(doc, layout_type)
    add_signature_section(doc, layout_type)

    os.makedirs(output_folder, exist_ok=True)
    docx = os.path.join(output_folder, f"daily_report_{doc_number:03}_{layout_type}.docx")
    pdf  = os.path.join(output_folder, f"daily_report_{doc_number:03}_{layout_type}.pdf")
    doc.save(docx)
    convert(docx, pdf)
    os.remove(docx)