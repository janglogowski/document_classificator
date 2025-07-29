from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx2pdf import convert
import os
import random
import datetime

#---lists used by the generator---#
fonts = ['Arial','Calibri']
inspectors = ['Anna Nowak', 'Jan Kowalski', 'Peter Schmidt', 'Laura Rossi', 'Carlos Garcia']
products = [
    'MC-540X', 'TR-200B', 'HF-390A', 'PL-601Z', 'DX-777T',
    'TX-820V', 'MX-450L', 'RX-310Z', 'VF-220D', 'GL-980S',
    'AL-115Q', 'KP-320E', 'BZ-660F', 'QN-770H', 'SL-430M',
    'ZR-205R', 'TY-350G', 'XK-610U', 'JD-700W', 'CN-150C',
    'VR-940T', 'MS-600P', 'LK-890B', 'FT-730X', 'NE-245A',
    'PW-515Y', 'RM-860N', 'WD-180S', 'KV-390K', 'CE-905L',
    'LP-555V', 'GH-770J', 'SB-140D', 'QP-660F', 'NU-440Z',
    'AZ-300T', 'XD-710R', 'RE-850C', 'MR-160H', 'TL-900X'
]
dimensions = ['Thickness', 'Width', 'Length', 'Hole Ø', 'Height', 'Depth', 'Inner Diameter']
components = [
    'Steel Sheet A36','Hex Bolts M12','Rubber Gasket 80mm','O-Ring NBR 60mm',
    'Bearing 6202 ZZ','Shaft 500mm','Plastic Rivets','Graphite Pad',
    'Battery Pack','Hinge Set','Power Switch','Wooden Pallet'
]
tools = ['Caliper', 'Micrometer', 'CMM', 'Laser Scanner', 'Depth Gauge', 'Measuring Tape']

column_synonyms = {
    "Product ID": ["Product Ref", "Item Code", "Article No."],
    "Component Name": ["Component", "Part Name"],
    "Dimension": ["Dim.", "Measurement"],
    "Nominal (mm)": ["Target", "Nominal"],
    "Measured (mm)": ["Observed", "Actual"],
    "Deviation": ["Diff", "Delta"],
    "Status (OK/NOK)": ["Pass/Fail"]
}

titles = [
    "Tolerance Check",
    "Measurement Summary Sheet",
    "Dimensional Log",
    "Inspection Results Summary",
    "Component Dimensional Check"
]

inspector_labels = ["Inspector", "Technician", "Supervisor"]

status_sets = [
    ("OK", "NOK"),
    ("PASS", "FAIL"),
    ("V", "X")
]

insp_intros = [
    "This report presents the dimensional measurements and inspection results.",
    "Below are the recorded measurements compared against nominal tolerances.",
    "The following data captures key dimensions and any deviations identified.",
    "Please review the inspection results for each component listed below.",
    "This section details the measured values, tolerances, and status flags.",
    "Use this examination summary to confirm component conformity.",
    "Entries include both pass/fail markers and deviation magnitudes.",
    "Ensure measurement methods align with calibration standards.",
    "Refer to the dimensional log for all component size readings.",
    "This summary of measurements supports metrology traceability.",
    "Review recorded tolerances against engineering specifications.",
    "The inspection register below highlights any out-of-tolerance parts.",
    "Check that all dimensions comply with ISO and company standards.",
    "Use this results summary to trigger any corrective actions.",
    "All measured values are timestamped for audit purposes.",
    "This data extract is prepared for quality-control sign-off."
]

insp_summaries = [
    "All dimensions within tolerance have been marked as OK.",
    "Components failing inspection require immediate review and corrective action.",
    "Refer to deviation column for any out-of-tolerance measurements.",
    "Overall inspection summary indicates acceptable quality levels.",
    "Please address any NOK items before proceeding to the next production stage.",
    "This assessment reflects the latest metrology results.",
    "Ensure all measuring tools were properly calibrated.",
    "Inspection notes have been logged for traceability.",
    "Confirm that pass rates meet the defined acceptance criteria.",
    "Archive this inspection summary for regulatory documentation.",
    "Flag any components requiring re-machining or adjustment.",
    "Review failed items against the corrective-action register.",
    "This closure memo confirms that dimensional checks are complete.",
    "Ensure status flags are updated in the quality management system.",
    "Cross-verify measurement data with CAD nominal values.",
    "Record any measurement anomalies for follow-up analysis."
]

def add_horizontal_line(paragraph):
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'auto')
    pBdr.append(bottom)
    pPr.append(pBdr)

def random_date():
    today = datetime.date.today()
    delta_days = random.randint(0, 730)
    return today - datetime.timedelta(days=delta_days)

def add_inspector_and_date(doc, inspectors, layout_type):
    label = random.choice(inspector_labels)
    inspector = random.choice(inspectors)
    date = random_date().strftime("%Y-%m-%d")

    if layout_type == 2:
        table = doc.add_table(rows=2, cols=2)
        table.style = 'Table Grid'
        table.cell(0, 0).text = label
        table.cell(0, 1).text = inspector
        table.cell(1, 0).text = "Inspection Date"
        table.cell(1, 1).text = date

        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(8.5)

        doc.add_paragraph()
    elif layout_type == 3:
        p = doc.add_paragraph(f"Inspection performed by {inspector} on {date}.")
        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        p.runs[0].font.size = Pt(10)
        p.runs[0].bold = True

def add_dimensional_header(doc, layout_type):
    if random.random() < 0.8:  
        p = doc.add_paragraph()
        run = p.add_run(random.choice(titles))
        run.bold = True
        run.font.size = Pt(14)
        if layout_type != 2:
            add_horizontal_line(p)

def add_dimensional_table(doc, products, components, dimensions, tools, layout_type):
    status_ok, status_nok = random.choice(status_sets)

    if layout_type == 2:
        headers = [
            "Product ID", "Component Name", "Dimension",
            "Nominal (mm)", "Measured (mm)", "Deviation", "Status (OK/NOK)"
        ]
        header_labels = [random.choice(column_synonyms.get(h, [h])) for h in headers]
        num_measurements = random.randint(4, 7)
        table = doc.add_table(rows=len(header_labels), cols=num_measurements + 1)
        table.style = 'Table Grid'

        ok_count = 0
        nok_count = 0

        for i, h in enumerate(header_labels):
            run = table.cell(i, 0).paragraphs[0].add_run(h)
            run.bold = True
            run.font.size = Pt(8.5)

        for col in range(1, num_measurements + 1):
            nominal = round(random.uniform(5.0, 100.0), 2)
            tol = 0.2
            measured = round(nominal + random.uniform(-tol, tol), 2)
            deviation = round(measured - nominal, 2)
            status = status_ok if abs(deviation) <= tol else status_nok

            if status == status_ok:
                ok_count += 1
            else:
                nok_count += 1

            values = [
                random.choice(products),
                random.choice(components),
                random.choice(dimensions),
                f"{nominal:.2f}",
                f"{measured:.2f}",
                f"{deviation:+.2f}",
                status
            ]

            for row_idx, val in enumerate(values):
                cell = table.cell(row_idx, col)
                cell.text = val
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.size = Pt(8.5)

        doc.add_paragraph()
        return ok_count, nok_count

    elif layout_type == 3:
        headers = [
            "Product ID", "Component Name", "Dimension",
            "Nominal (mm)", "Measured (mm)", "Deviation", "Status (V/X)"
        ]
    else:
        headers = [
            "Product ID", "Component", "Dimensions",
            "Nominal (mm)", "Tolerance (mm)", "Measured (mm)",
            "Deviation", "Tool", "Status (Good/Bad)"
        ]

    header_labels = [random.choice(column_synonyms.get(h, [h])) for h in headers]
    table = doc.add_table(rows=1, cols=len(header_labels))
    table.style = 'Table Grid'

    for i, h in enumerate(header_labels):
        run = table.cell(0, i).paragraphs[0].add_run(h)
        run.bold = True

    ok_count = 0
    nok_count = 0
    for _ in range(random.randint(10, 18)):
        nominal = round(random.uniform(5.0, 100.0), 2)
        tol = 0.2
        measured = round(nominal + random.uniform(-tol, tol), 2)
        deviation = round(measured - nominal, 2)
        status = status_ok if abs(deviation) <= tol else status_nok

        if status == status_ok:
            ok_count += 1
        else:
            nok_count += 1

        row = table.add_row().cells
        base_values = [
            random.choice(products),
            random.choice(components),
            random.choice(dimensions),
            f"{nominal:.2f}",
        ]

        if layout_type == 1:
            base_values += [
                "±0.20",
                f"{measured:.2f}",
                f"{deviation:+.2f}",
                random.choice(tools),
                status
            ]
        else:
            base_values += [
                f"{measured:.2f}",
                f"{deviation:+.2f}",
                status
            ]

        for i in range(len(base_values)):
            if i < len(row):
                row[i].text = base_values[i]

    doc.add_paragraph()
    return ok_count, nok_count

def add_summary_section(doc, ok_count, nok_count):
    formats = [
        f"Summary – Passed: {ok_count}, Failed: {nok_count}",
        f"Final Count: {ok_count} OK, {nok_count} NOK",
        f"Inspected Units – V: {ok_count} / X: {nok_count}",
    ]
    p = doc.add_paragraph()
    run = p.add_run(random.choice(formats))
    run.bold = True
    run.font.size = Pt(10)

def add_calibration_log(doc):
    doc.add_paragraph("Instrument Calibration Log:")
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr = ["Instrument", "Serial No.", "Last Calibration Date"]
    for i, h in enumerate(hdr):
        cell = table.cell(0, i)
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True

    instruments = ["Caliper", "Micrometer", "CMM", "Laser Scanner"]
    for _ in range(random.randint(2,4)):
        inst = random.choice(instruments)
        ser  = f"{random.randint(10000,99999)}"
        date = (datetime.date.today() - datetime.timedelta(days=random.randint(30,360))).strftime("%Y-%m-%d")
        row = table.add_row().cells
        row[0].text, row[1].text, row[2].text = inst, ser, date
    doc.add_paragraph("")

def add_environmental_conditions(doc):
    temp = f"{random.uniform(20.0,25.0):.1f} °C"
    hum  = f"{random.uniform(30.0,60.0):.0f} %"
    para = doc.add_paragraph()
    para.add_run("Environmental Conditions: ").bold = True
    para.add_run(f"Temperature: {temp}, Humidity: {hum}")

def generate_dimensional_inspection_report(doc_number: int, output_folder: str):
    doc = Document()
    layout_type = random.randint(1, 3)
    #layout_type = 3

    font = random.choice(fonts)
    style = doc.styles['Normal']
    style.font.name = font
    style.font.size = Pt(9)

    add_dimensional_header(doc, layout_type)
    add_inspector_and_date(doc, inspectors, layout_type)

    if layout_type == 2:
        if random.random() < 0.7:
            intro_sentences = random.sample(insp_intros, k=random.randint(3,6))
            if intro_sentences:
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after  = Pt(0)
                for i, sent in enumerate(intro_sentences):
                    run = p.add_run(sent)
                    if i < len(intro_sentences) - 1:
                        run.add_text(" ")
            doc.add_paragraph("")

    if layout_type == 3:
        if random.random() < 0.7:
            intro_sentences = random.sample(insp_intros, k=random.randint(4,6))
            if intro_sentences:
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after  = Pt(0)
                for i, sent in enumerate(intro_sentences):
                    run = p.add_run(sent)
                    if i < len(intro_sentences) - 1:
                        run.add_text(" ")
            add_environmental_conditions(doc)

    ok, nok = add_dimensional_table(doc, products, components, dimensions, tools, layout_type)

    if layout_type == 2:
        intro_sentences = random.sample(insp_summaries, k=random.randint(3,4))
        if intro_sentences:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after  = Pt(0)
            for i, sent in enumerate(intro_sentences):
                run = p.add_run(sent)
                if i < len(intro_sentences) - 1:
                    run.add_text(" ")

        line = doc.add_paragraph()
        add_horizontal_line(line)
        add_calibration_log(doc)

    if layout_type == 1:
        add_summary_section(doc, ok, nok)

    #---save docx, convert to pdf, remove docx---#
    os.makedirs(output_folder, exist_ok=True)
    docx_path = os.path.join(output_folder, f"dim_inspection_report_{doc_number:03}_{layout_type}.docx")
    pdf_path  = os.path.join(output_folder, f"dim_inspection_report_{doc_number:03}_{layout_type}.pdf")
    doc.save(docx_path)
    convert(docx_path, pdf_path)
    os.remove(docx_path)