from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx2pdf import convert
import random
import os
import datetime

#---lists used by the generator---#
supervisors = ['Anna Nowak', 'Jan Kowalski', 'Peter Schmidt', 'Laura Rossi', 'Carlos Garcia']
equipment_names = [
    "CNC Milling Machine", "Hydraulic Press", "Laser Cutter", "Lathe", "Plasma Cutter",
    "Assembly Robot", "Paint Booth", "Packaging Line", "Conveyor Belt"]
maintenance_types = ["Preventive", "Corrective", "Inspection"]
remarks_pool = [
    "Changed oil and filters.", "No issues found.", "Replaced coolant.",
    "Sensor recalibrated.", "Calibration check OK.", "Replaced fuse (5A).",
    "Worn gasket replaced.", "Tightened loose bolts.", "Refilled oil (HLP 46).",
    "Li-Ion battery pack serviced.", "Alignment of hinges adjusted.",
    "Switch contacts cleaned.", "Wooden pallet checked.",""]
locations = ["Plant 3A", "Plant 2B", "Plant 1C"]
fonts = ['Arial', 'Calibri', 'Aptos']

titles = [
    "Equipment Record Report", "Maintenance Checklist Report", "Machine Log Summary"
]
signature_roles = [
    ("Approved by:", "Serviced by:"),
    ("Signed off:", "Performed by:"),
    ("Authorized by:", "Operator:"),
]
column_synonyms = {
    "Equipment ID": ["Machine ID", "Unit Code"],
    "Equipment Name": ["Machine", "Equipment"],
    "Maintenance Type": ["Type", "Service Type"],
    "Performed By": ["Operator", "Technician"],
    "Downtime (hrs)": ["Downtime", "Duration"],
    "Last Maintenance": ["Previous Service", "Last Check"],
    "Location": ["Site", "Area"],
    "Remarks": ["Notes", "Comments"]
}

maint_intros = [
    "This log records all maintenance activities performed on the equipment.",
    "Below is the summary of service actions and downtimes for each machine.",
    "The following entries detail preventive and corrective maintenance tasks.",
    "Please review this log for recent service dates, types, and remarks.",
    "This section captures equipment status and any maintenance observations.",
    "Use this record to plan upcoming upkeep and repairs.",
    "Entries include work orders, technician assignments, and part changes.",
    "Ensure all safety checks were completed during servicing.",
    "Refer to the maintenance register for detailed service histories.",
    "This service summary supports the reliability engineering review.",
    "Check that all downtime events are properly categorized and logged.",
    "The equipment log below includes fault codes and corrective notes.",
    "Use this summary to forecast parts replacement and resource needs.",
    "All technician comments are recorded for maintenance trend analysis.",
    "This maintenance extract is prepared for compliance audit trails.",
    "Confirm that service intervals follow the preventive schedule."
]

maint_summaries = [
    "All maintenance tasks have been completed as per schedule.",
    "No critical faults were found during the latest inspection.",
    "Refer to remarks for any follow-up actions or parts replacements.",
    "Ensure next preventive service is scheduled according to plan.",
    "Overall equipment condition is satisfactory post-maintenance.",
    "Service summaries have been forwarded to the engineering team.",
    "Record any spare parts usage for inventory adjustment.",
    "Maintenance notes are archived for compliance audits.",
    "Confirm that all corrective actions were properly closed out.",
    "This log summary supports the asset-management dashboard.",
    "Check remaining life-cycles of key components from this report.",
    "Flag any recurring issues for root-cause investigation.",
    "Archive this summary in the CMMS for future reference.",
    "Ensure that each service entry has the required approvals.",
    "Use this closure note to update the maintenance KPI tracker.",
    "All maintenance durations are recorded for performance metrics."
]

def get_synonym(key):
    return random.choice(column_synonyms.get(key,[key]))

def get_random_date_within_2_years():
    today = datetime.date.today()
    delta_days = random.randint(0, 730)
    return today - datetime.timedelta(days=delta_days)

def assign_unique_equipment_entries(min_rows=4, max_rows=10):
    used_ids = set()
    equipment_log = []
    num_rows = random.randint(min_rows, max_rows)

    for _ in range(num_rows):
        while True:
            equipment_id = f"MC-2{random.randint(1, 99):02}"
            if equipment_id not in used_ids:
                used_ids.add(equipment_id)
                equipment_name = random.choice(equipment_names)
                equipment_log.append((equipment_id, equipment_name))
                break
    return equipment_log

def add_horizontal_line(paragraph):
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'auto')
    pBdr.append(bottom)
    pPr.append(pBdr)

def add_title(doc, layout_type):
    if random.random() < 0.2: 
        return

    p = doc.add_paragraph()
    title = random.choice(titles)

    if layout_type == 2:
        p.paragraph_format.tab_stops.add_tab_stop(Inches(6), alignment=WD_PARAGRAPH_ALIGNMENT.RIGHT)
        run = p.add_run(f"{title}\t{get_random_date_within_2_years().strftime('%Y-%m-%d')}")
    else:
        p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        suffix = f" #{random.randint(1000000,9999999)}" if layout_type == 1 else ""
        run = p.add_run(title + suffix)

    run.font.size = Pt(11)
    run.bold = True
    add_horizontal_line(p)

    if layout_type != 2:
        para = doc.add_paragraph()
        label = "Date:" if layout_type == 1 else "Performed On:"
        run = para.add_run(f"{label} {get_random_date_within_2_years().strftime('%Y-%m-%d')}")
        run.font.size = Pt(9)

def add_maintenance_log_table(doc, layout_type, equipment_log):
    base_headers = [
        "Equipment ID", "Equipment Name", "Maintenance Type",
        "Performed By", "Downtime (hrs)"
    ]
    if layout_type == 2:
        base_headers.append("Last Maintenance")
    if layout_type != 3:
        base_headers.append("Location")
    base_headers.append("Remarks")

    headers = [random.choice(column_synonyms.get(h, [h])) for h in base_headers]

    if layout_type == 3:
        n = min(len(equipment_log), 6)
        entries_a = random.sample(equipment_log, n)
        entries_b = random.sample(equipment_log, n)

        for section_label, entries in [("Section A:", entries_a), ("Section B:", entries_b)]:
            doc.add_paragraph(section_label).runs[0].bold = True

            table = doc.add_table(rows=len(headers), cols=n + 1)
            table.style = 'Table Grid'

            for row_idx, h in enumerate(headers):
                cell = table.cell(row_idx, 0)
                run = cell.paragraphs[0].add_run(h)
                run.bold = True

            for col_idx, (equipment_id, equipment_name) in enumerate(entries, start=1):
                values = [
                    equipment_id,
                    equipment_name,
                    random.choice(maintenance_types),
                    random.choice(supervisors),
                    f"{random.uniform(1.5, 4.0):.1f}",
                    random.choice(remarks_pool)
                ]
                for row_idx, val in enumerate(values):
                    table.cell(row_idx, col_idx).text = val

            doc.add_paragraph("")

        return

    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    for i, h in enumerate(headers):
        run = table.cell(0, i).paragraphs[0].add_run(h)
        run.bold = True

    for equipment_id, equipment_name in equipment_log:
        row = table.add_row().cells
        values = [
            equipment_id,
            equipment_name,
            random.choice(maintenance_types),
            random.choice(supervisors),
            f"{random.uniform(1.5, 4.0):.1f}"
        ]
        if layout_type == 2:
            values.append(get_random_date_within_2_years().strftime('%Y-%m-%d'))
        else:
            values.append(random.choice(locations))
        values.append(random.choice(remarks_pool))

        for i, val in enumerate(values):
            row[i].text = val

    doc.add_paragraph("")

def add_performance_section(doc):
    tbl = doc.add_table(rows=1,cols=3,style='Table Grid')
    hdr = ["Uptime (%)","Error Count",get_synonym("Remarks")]
    for i,h in enumerate(hdr):
        c=tbl.cell(0,i); c.text=h; c.paragraphs[0].runs[0].bold=True
    r=tbl.add_row().cells
    r[0].text = f"{random.uniform(90,99):.2f}%"
    r[1].text = str(random.randint(0,5))
    r[2].text = random.choice(["","Replaced filter","Bolt tightened"])
    doc.add_paragraph()

def add_checklist_section(doc):
    tbl=doc.add_table(rows=3,cols=2,style='Table Grid')
    pts=["Lubrication Checked","Calibration Verified","Emergency Stop Tested"]
    for i,l in enumerate(pts):
        tbl.cell(i,0).text=l
        tbl.cell(i,1).text=random.choice(["Yes","No","N/A"])
    doc.add_paragraph()

def add_signature(doc, layout_type):
    roles = random.choice(signature_roles)
    people = [random.choice(supervisors), random.choice(supervisors)]

    if layout_type == 1:
        doc.add_paragraph(f"{roles[0]} _______________________          {roles[1]} _______________________")
    else:
        table = doc.add_table(rows=2, cols=2)
        table.style = 'Table Grid'
        for i in range(2):
            table.cell(i, 0).text = roles[i]
            table.cell(i, 1).text = people[i]

def add_spare_parts_section(doc):
    parts_pool = [
        "Oil Filter", "Hydraulic Hose", "Sealing Gasket", "Sensor Module",
        "Fuse 5A", "Li-Ion Battery Pack", "Hinge Set", "Switch Assembly"
    ]

    n = random.randint(1, 4)
    table = doc.add_paragraph("Spare Parts Used:").insert_paragraph_before().add_run()._r
    tbl = doc.add_table(rows=1, cols=3)
    tbl.style = 'Table Grid'
    hdr = ["Part Name", "Part No.", "Qty Used"]
    for i, h in enumerate(hdr):
        cell = tbl.cell(0, i)
        run = cell.paragraphs[0].add_run(h)
        run.bold = True

    for _ in range(n):
        part = random.choice(parts_pool)
        part_no = f"P-{random.randint(1000,9999)}"
        qty = random.randint(1, 5)
        row = tbl.add_row().cells
        row[0].text = part
        row[1].text = part_no
        row[2].text = str(qty)

    doc.add_paragraph("")

def generate_maintenance_log(doc_number: int, output_folder: str):
    doc = Document()
    layout_type = random.randint(1, 3)
    # layout_type = 3

    font_name = random.choice(fonts)
    font_size = Pt(8)
    style = doc.styles['Normal']
    style.font.name = font_name
    style.font.size = font_size

    equipment_log = assign_unique_equipment_entries()

    add_title(doc, layout_type)

    if layout_type == 2:
        intro_sentences = random.sample(maint_intros, k=random.randint(4,6))
        if intro_sentences:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after  = Pt(0)
            for i, sent in enumerate(intro_sentences):
                run = p.add_run(sent)
                if i < len(intro_sentences) - 1:
                    run.add_text(" ")
        add_spare_parts_section(doc)
        add_maintenance_log_table(doc, layout_type, equipment_log)
        add_signature(doc, layout_type)

    if layout_type == 3:
        add_signature(doc, layout_type)
        doc.add_paragraph("")
        intro_sentences = random.sample(maint_intros, k=random.randint(5,7))
        if intro_sentences:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after  = Pt(0)
            for i, sent in enumerate(intro_sentences):
                run = p.add_run(sent)
                if i < len(intro_sentences) - 1:
                    run.add_text(" ")
        doc.add_paragraph("")

        add_maintenance_log_table(doc, layout_type, equipment_log)

        if random.random() < 0.6:
            intro_sentences = random.sample(maint_summaries, k=random.randint(2,7))
            if intro_sentences:
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after  = Pt(0)
                for i, sent in enumerate(intro_sentences):
                    run = p.add_run(sent)
                    if i < len(intro_sentences) - 1:
                        run.add_text(" ")
            doc.add_paragraph("")

        add_performance_section(doc)
        
    if layout_type == 1:
        add_maintenance_log_table(doc, layout_type, equipment_log)
        intro_sentences = random.sample(maint_summaries, k=random.randint(4,6))
        if intro_sentences:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after  = Pt(0)
            for i, sent in enumerate(intro_sentences):
                run = p.add_run(sent)
                if i < len(intro_sentences) - 1:
                    run.add_text(" ")
        doc.add_paragraph("")
        add_checklist_section(doc)
        add_signature(doc, layout_type)

    #---save docx, convert to pdf, remove docx---#
    os.makedirs(output_folder, exist_ok=True)
    docx_path = os.path.join(output_folder, f"maintenance_log_{doc_number:03}_{layout_type}.docx")
    pdf_path  = os.path.join(output_folder, f"maintenance_log_{doc_number:03}_{layout_type}.pdf")
    doc.save(docx_path)
    convert(docx_path, pdf_path)
    os.remove(docx_path)