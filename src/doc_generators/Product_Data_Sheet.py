from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx2pdf import convert
import os
import random
import datetime

#---lists used by the generator---#
fonts = ['Arial', 'Calibri']
locations = ["Plant 3A", "Plant 2B", "Plant 1C"]
supervisors = ['Anna Nowak', 'Jan Kowalski', 'Peter Schmidt', 'Laura Rossi', 'Carlos Garcia']

products = [
    'MC-540X', 'TR-200B', 'HF-390A', 'PL-601Z', 'DX-777T',
    'TX-820V', 'MX-450L', 'RX-310Z', 'VF-220D', 'GL-980S',
    'AL-115Q', 'KP-320E', 'BZ-660F', 'QN-770H', 'SL-430M',
    'ZR-205R', 'TY-350G', 'XK-610U', 'JD-700W', 'CN-150C',
    'VR-940T', 'MS-600P', 'LK-890B', 'FT-730X', 'NE-245A',
    'PW-515Y', 'RM-860N', 'WD-180S', 'KV-390K', 'CE-905L',
    'LP-555V', 'GH-770J', 'SB-140D', 'QP-660F', 'NU-440Z',
    'AZ-300T', 'XD-710R', 'RE-850C', 'MR-160H', 'TL-900X']

item_to_material = {
    'Steel Sheet A36': "Carbon Steel A36",
    'Hex Bolts M12': "Zinc-Plated Steel",
    'Rubber Gasket 80mm': "Nitrile Rubber (NBR)",
    'O-Ring NBR 60mm': "Nitrile Rubber (NBR)",
    'Bearing 6202 ZZ': "Stainless Steel 304",
    'Graphite Pad': "Graphite",
    'Teflon Tape Roll': "PTFE",
    'Battery Pack': "Lithium-Ion Cell Pack",
    'Hinge Set': "Stainless Steel 304",
    'Power Switch': "Polycarbonate + Copper",
    'Wooden Pallet': "Treated Pine Wood",
    'Digital Display Unit': "ABS Plastic",
    'Cable Tie Pack (100)': "Nitrile Rubber (NBR)",
    'Gasket Sheet A4': "Nitrile Rubber (NBR)",
    'O-Ring NBR 60mm': "Nitrile Rubber (NBR)",
    'Fuse 5A': "Ceramic Oxide",
    'LED Light Strip': "Glass-Filled Nylon",
    'Cooling Gel Pack': "Silicone Rubber",
    'Graphite Pad': "Graphite Pad",
    'Support Foot Steel': "Stainless Steel 304",
    'Wiring Loom 1m': "PVC (Rigid)",
    'Insulated Tube 25mm': "Thermoplastic Polyurethane (TPU)",
    'Pressure Valve': "Brass CZ121",
    'Capacitor 450V': "Aluminum 6061-T6",
    'Thermal Fuse': "Ceramic Oxide",
    'Nut M6': "Zinc-Plated Steel",
    'Connector 2P': "Glass-Filled Nylon",
    'Heat Sink ALU': "Aluminum 6061-T6",
    'Teflon Tape Roll': "Teflon Tape Roll"}

synonym_titles = ["Component Overview", "Unit Data Specification", "Reference Configuration","Part Configuration Log"]
section_order = ["param_first", "mat_first", "only_param", "only_mat"]

column_synonyms = {
    "Component Name": ["Item", "Label", "Part", "Subcomponent"],
    "Material": ["Raw Material", "Composition", "Base"]}

pds_intros = [
    "This data sheet provides key specifications and material details.",
    "Below are the technical parameters and construction materials for the unit.",
    "The following information outlines performance characteristics and design data.",
    "Please review the specifications and material composition listed below.",
    "This section details the configuration and parameter set for the product.",
    "Use this sheet to confirm engineering requirements and tolerances.",
    "Entries include both mechanical and electrical specifications.",
    "Ensure all referenced standards are up to date.",
    "Refer to this configuration summary for unit-designation mapping.",
    "The parameter register below includes operating limits and ratings.",
    "Check that all material specs align with supplier certifications.",
    "This technical overview supports the product-release checklist.",
    "Use this spec sheet to validate assembly instructions.",
    "All data entries are traceable to design revision history.",
    "Confirm that performance ranges comply with project requirements.",
    "This report extract is prepared for design-verification audits."]

pds_summaries = [
    "All specifications meet the design requirements and industry standards.",
    "Material choices and process data have been verified for compliance.",
    "Refer to parameter table for operating ranges and tolerances.",
    "Ensure that the listed materials are approved for the application.",
    "Overall configuration is within defined design and safety margins.",
    "This summary reflects the latest revision of the product data.",
    "Document control numbers are included for traceability.",
    "Data sheet has been peer-reviewed by the engineering team.",
    "Confirm that all test conditions are accurately represented.",
    "Archive this summary for future design-change management.",
    "Check that all section headers follow the product template.",
    "Ensure glossary terms match the engineering nomenclature.",
    "Use this summary to cross-check with BOM and inspection reports.",
    "Flag any missing parameters for urgent specification updates.",
    "This closure note indicates the data sheet is ready for release.",
    "All summary comments have been recorded in the revision log."]

parameter_pool = {
    "Unit Type": ["Hydraulic Power Unit", "Pneumatic Control Unit", "Cooling Circulation System"],
    "Design Pressure": ["180 bar", "210 bar", "250 bar", "300 bar"],
    "Flow Rate": ["30 L/min", "48 L/min", "60 L/min", "75 L/min"],
    "Reservoir Capacity": ["80 L", "120 L", "160 L"],
    "Motor Power": ["5.5 kW", "7.5 kW", "11 kW"],
    "Voltage": ["230 V / 50 Hz", "400 V / 50 Hz", "480 V / 60 Hz"],
    "Pump Type": ["Vane Pump", "Gear Pump", "Piston Pump"],
    "Control Valve": ["Directional 4/3", "Proportional 4/2", "Manual Override 3/2"],
    "Filtration": ["10 μm return, 25 μm suction", "5 μm inline, 10 μm return"],
    "Cooling": ["Air-cooled oil radiator", "Water-cooled exchanger"],
    "Oil Type": ["HLP 32", "HLP 46", "HLP 68"],
    "Working Temp. Range": ["10°C to 55°C", "5°C to 45°C", "-10°C to 50°C"],
    "Frame Material": ["Powder-coated Steel", "Aluminum Frame", "Stainless Steel Frame"],
    "Protection Class": ["IP54", "IP65", "IP67"],
    "Noise Level": ["≤ 68 dB(A)", "≤ 70 dB(A)", "≤ 72 dB(A)"],
    "Total Weight": ["180 kg", "230 kg", "250 kg"],
    "Dimensions (LxWxH)": ["1000 x 650 x 1000 mm", "1200 x 700 x 1100 mm"],
    "Mounting": ["Skid base with vibration pads", "Wall-mounted", "Rack frame"],
    "Battery Capacity": ["2 Ah", "5 Ah", "10 Ah"],
    "Charging Time": ["2 h", "4 h", "6 h"],
    "Service Interval": ["500 h", "1000 h", "12 months"]}

parameter_synonyms = {
    "Unit Type": ["Unit Type", "Model Type", "Configuration Type"],
    "Design Pressure": ["Design Pressure", "Operating Pressure", "Rated Pressure"],
    "Flow Rate": ["Flow Rate", "Throughput", "Fluid Flow"],
    "Voltage": ["Voltage", "Operating Voltage", "Supply Voltage"]}

def generate_parameter_subset(layout_type):
    essential_keys = ["Unit Type", "Design Pressure", "Flow Rate", "Voltage"]
    optional_keys = [k for k in parameter_pool if k not in essential_keys]

    if layout_type == 1:
        n = random.randint(12, len(optional_keys))
    elif layout_type == 2:
        n = random.randint(3, 6)
    else:
        n = random.randint(5, 8)

    selected_keys = essential_keys + random.sample(optional_keys, n)

    result = []
    for key in selected_keys:
        header_choices = parameter_synonyms.get(key, [key])
        header = random.choice(header_choices)
        value = random.choice(parameter_pool[key])
        result.append((header, value))
    return result

def random_date():
    today = datetime.date.today()
    delta_days = random.randint(0, 730)
    return today - datetime.timedelta(days=delta_days)

def add_horizontal_line(paragraph):
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'auto')
    pBdr.append(bottom)
    pPr.append(pBdr)

def add_product_header(doc, layout_type):
    if random.random() < 0.1:
        return 
    
    title = random.choice(["Product Overview","Product Data Sheet","Unit Configuration", "Component Overview", "Unit Data Specification"])

    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT if layout_type == 1 else WD_PARAGRAPH_ALIGNMENT.LEFT
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(14)

    if layout_type == 1:
        add_horizontal_line(p)

def add_meta_info(doc, layout_type):
    date = random_date().strftime('%Y-%m-%d')
    author = random.choice(supervisors)
    approver = random.choice(supervisors)
    location = random.choice(locations)
    
    doc_no_label = random.choice(["Document No.", "Report No.", "Internal No."])
    if random.random() < 0.5:
        doc_no = f"{random.randint(1000000, 9999999)}"
    else:
        doc_no = f"PR-{random.randint(100, 999)}"

    product = random.choice(products)

    if layout_type == 2:
        table = doc.add_table(rows=2, cols=2)
        table.style = 'Table Grid'
        table.cell(0, 0).text = f"Product ID:"
        table.cell(0, 1).text = product
        table.cell(1, 0).text = doc_no_label + ":"
        table.cell(1, 1).text = doc_no
        doc.add_paragraph(f"Prepared by: {author} | Approved by: {approver} | Date: {date}")

    elif layout_type == 3:
        p = doc.add_paragraph()
        p.add_run(f"{doc_no_label}: {doc_no}\t\tProduct ID: {product}\t\tDate: {date}")
        return author, approver

    else:
        table = doc.add_table(rows=4, cols=2)
        table.style = 'Table Grid'
        table.cell(0, 0).text = "Project:"
        table.cell(0, 1).text = product
        table.cell(1, 0).text = doc_no_label + ":"
        table.cell(1, 1).text = doc_no
        table.cell(2, 0).text = "Prepared By:"
        table.cell(2, 1).text = author
        table.cell(3, 0).text = "Approved By:"
        table.cell(3, 1).text = approver
        doc.add_paragraph(f"Location: {location} | Date: {date}")

def add_parameter_table(doc, layout_type):
    parameters = generate_parameter_subset(layout_type)
    if isinstance(parameters, dict):
        items = list(parameters.items())
    else:
        items = parameters 

    num = len(items)
    rows = (num + 1) // 2
    table = doc.add_table(rows=rows, cols=4)
    table.style = 'Table Grid'

    idx = 0
    for r in range(rows):
        for c in (0, 2):
            if idx < num:
                header, value = items[idx]
                cell_h = table.cell(r, c)
                cell_h.text = header
                for run in cell_h.paragraphs[0].runs:
                    run.bold = True
                table.cell(r, c + 1).text = value
                idx += 1

    doc.add_paragraph("") 

def add_material_table(doc, layout_type):
    specs = ["DIN EN 10025", "ISO 2768", "Class A tolerance", "IP67 rated","Heat-resistant", "Electrical grade", "UL-listed", "Anticorrosive","RoHS compliant", "CE certified"]

    layout_variant = random.choice(section_order)
    if layout_variant == "only_param":
        return 

    col1 = random.choice(column_synonyms["Component Name"])
    col2 = random.choice(column_synonyms["Material"])

    if layout_type == 3:
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        header_cell = table.cell(0, 0).merge(table.cell(0, 2))
        header_cell.text = "Component Materials Overview:"
        header_cell.paragraphs[0].runs[0].bold = True

        row = table.add_row().cells
        row[0].text = col1
        row[1].text = col2
        row[2].text = "Specification"

        sampled_items = random.sample(list(item_to_material.items()), k=random.randint(3, 8))
        for component, material in sampled_items:
            row = table.add_row().cells
            row[0].text = component
            row[1].text = material.replace("Material", random.choice(["Base", "Spec", "Core"]))
            row[2].text = random.choice(specs)

    else:
        doc.add_paragraph("Material of Construction:")
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        table.cell(0, 0).text = col1
        table.cell(0, 1).text = col2

        sampled_items = random.sample(list(item_to_material.items()), k=random.randint(3, 12))
        for component, material in sampled_items:
            row = table.add_row().cells
            row[0].text = component
            row[1].text = material

    doc.add_paragraph("")

def generate_product_data_sheet(doc_number: int, output_folder: str):
    doc = Document()

    layout_type = random.randint(1, 3)
    #layout_type = 3
    font = random.choice(fonts)
    style = doc.styles['Normal']
    style.font.name = font
    style.font.size = Pt(9)

    add_product_header(doc, layout_type)

    if layout_type == 2:
        intro_sentences = random.sample(pds_intros, k=random.randint(5,7))
        if intro_sentences:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after  = Pt(0)
            for i, sent in enumerate(intro_sentences):
                run = p.add_run(sent)
                if i < len(intro_sentences) - 1:
                    run.add_text(" ")
        doc.add_paragraph("")

    if layout_type == 3:
        author, approver = add_meta_info(doc, layout_type)
    else:
        add_meta_info(doc, layout_type)

    if layout_type == 1:
        if random.random() < 0.7:
            intro_sentences = random.sample(pds_intros, k=random.randint(4,7))
            if intro_sentences:
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after  = Pt(0)
                for i, sent in enumerate(intro_sentences):
                    run = p.add_run(sent)
                    if i < len(intro_sentences) - 1:
                        run.add_text(" ")
            doc.add_paragraph("")

    add_parameter_table(doc, layout_type)

    if layout_type == 2:
        intro_sentences = random.sample(pds_summaries, k=random.randint(5,7))
        if intro_sentences:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after  = Pt(0)
            for i, sent in enumerate(intro_sentences):
                run = p.add_run(sent)
                if i < len(intro_sentences) - 1:
                    run.add_text(" ")
        doc.add_paragraph("")

    if layout_type == 1:
        if random.random() < 0.4:
            intro_sentences = random.sample(pds_summaries, k=random.randint(4,7))
            if intro_sentences:
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after  = Pt(0)
                for i, sent in enumerate(intro_sentences):
                    run = p.add_run(sent)
                    if i < len(intro_sentences) - 1:
                        run.add_text(" ")
            doc.add_paragraph("")

    add_material_table(doc, layout_type)

    if layout_type == 3:
        footer = doc.add_paragraph()
        footer.paragraph_format.tab_stops.add_tab_stop(Pt(250), WD_PARAGRAPH_ALIGNMENT.RIGHT)
        run1 = footer.add_run("Prepared by: " + author)
        run1.bold = True
        run2 = footer.add_run("\t\tApproved by: " + approver)
        run2.bold = True

    #---save docx, convert to pdf, remove docx---#
    os.makedirs(output_folder, exist_ok=True)
    docx_path = os.path.join(output_folder, f"product_data_sheet_{doc_number:03}_{layout_type}.docx")
    pdf_path  = os.path.join(output_folder, f"product_data_sheet_{doc_number:03}_{layout_type}.pdf")
    doc.save(docx_path)
    convert(docx_path, pdf_path)
    os.remove(docx_path)